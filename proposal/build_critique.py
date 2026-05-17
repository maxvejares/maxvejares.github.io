"""Build the website critique .docx."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Base style
style = doc.styles['Normal']
style.font.name = 'Garamond'
style.font.size = Pt(11)
pf = style.paragraph_format
pf.space_before = Pt(0)
pf.space_after = Pt(0)
pf.line_spacing = 1.15
pf.first_line_indent = Inches(0)

def H(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    r.font.name = 'Garamond'
    r.font.bold = True
    r.font.size = Pt(14 if level == 1 else 12)
    return p

def P(text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.first_line_indent = Inches(0)
    r = p.add_run(text)
    r.font.name = 'Garamond'
    r.font.size = Pt(11)
    if italic:
        r.font.italic = True
    return p

# Title
title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(0)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = title.add_run("Website Critique")
r.font.name = 'Garamond'; r.font.size = Pt(18); r.font.bold = True

sub = doc.add_paragraph()
sub.paragraph_format.space_before = Pt(0); sub.paragraph_format.space_after = Pt(12)
r = sub.add_run("maxvejares.com (working folder, May 2026)")
r.font.name = 'Garamond'; r.font.size = Pt(11); r.font.italic = True

# OVERALL READ
H("Overall")
P("The site reads as a competent academic homepage with a coherent editorial aesthetic: warm paper background, oxblood accent, Spectral serif body, hairline rules. The information architecture is conventional and legible (About, Research, Policy, Data, Atlas, Teaching, Media, CV), and the bilingual mirror is a real asset. The two structural problems are that the navigation is injected by JavaScript (so the page has no header for crawlers, screen readers, or users with JS disabled), and that the homepage hero photo and the closing Cicarelli painting are hot-linked from a Squarespace CDN that you do not control. The first is a discoverability and accessibility problem; the second is a single point of failure that can blank both visual anchors of the homepage without warning.")
P("The prose across pages varies in directness. The Landholding working paper card is the strongest: puzzle, claim, design, finding, in that order. The other working papers and the book project open with rhetorical questions and arrive at the argument too late. The home page leads with affiliation rather than with what you actually study, which is the standard but weaker opening for an editorial site. The two recommendations below address those issues without changing the visual system.")

# CRITICAL
H("Critical fixes before publishing")

H("1. Bake the navigation into HTML, not JavaScript.", level=2)
P("The header and footer are constructed inside js/nav.js and appended to the DOM at load time. A reader with JS disabled, a search engine crawler, or a screen reader running before scripts finish sees a page with no navigation at all. The current approach also produces a visible flash on slow connections. Move the header markup into each .html file directly. If you want to keep the page list in one place, run a small build step (a five-line Python script) that injects the same HTML at build time, rather than at runtime in the browser.")

H("2. Self-host the photo and the painting.", level=2)
P("Both images are pulled from images.squarespace-cdn.com under an account ID (67e9a39b...) that presumably belongs to your prior Squarespace site. If that account is closed or the asset is renamed, the homepage breaks silently. Download the photo and the Cicarelli to img/ and reference them locally. The img/ folder already exists.")

H("3. Consolidate the four CV files.", level=2)
P("The folder contains CV-Maximiliano-Vejares.pdf, CV-Maximiliano-Vejares-NEW.pdf, CV-Maximiliano-Vejares-ES.pdf, and CV-Maximiliano-Vejares-ES-NEW.pdf. cv.html links to the -NEW English version; the Spanish equivalent is presumably broken or unverified. Pick one English and one Spanish file, rename without -NEW, and delete the rest. Then drop the in-page iframe: a download button plus a PDF preview is redundant, the iframe renders poorly on mobile, and many readers want the file, not the on-page reader.")

H("4. Remove the development leftover.", level=2)
P("font-compare.html is in the root and will deploy with the site. Delete it.")

H("5. Translate Spanish working-paper titles, or hold them in English consistently.", level=2)
P("es/research.html keeps English titles for Landholding, State Capacity and Industrial Policy, the Observatory, and Taming Nature, but translates the abstracts. A Spanish reader expects either a translated title with the English title in parentheses, or a consistent rule that working titles remain in English until published. The current mixture reads as incomplete translation.")

# CONTENT
H("Content recommendations")

H("Homepage: lead with what you study, not where you sit.", level=2)
P("The current opening (\"I am a Research Fellow at the Net Zero Industrial Policy Lab...\") is informationally correct but trades the most valuable real estate on the site for an affiliation that already appears in the footer. The same logic that governs paper abstracts applies here: lead with the question and the claim. A version closer to your strengths would read: \"I study how state authority took shape in nineteenth-century Latin America, and what that history reveals about the politics of industrial policy today. My research combines archival data with subnational panel methods to explain why state capacity varied so sharply across regions of the same country.\" Then a second paragraph for affiliation, then the contact line. The third paragraph (Chile MFA, Chicago, WZB, UNDP) is useful but mid-career enough to drop to a short \"prior positions\" line.")

H("Research page: tighten the working-paper cards.", level=2)
P("The Landholding card already works: it opens with a puzzle, names the mechanism, states the design, and gives the finding (\"the franchise expanded significantly more in hacienda-heavy departments after the 1874 electoral reform\"). The other three could match this structure.")
P("State Capacity and Industrial Policy should lead with its claim, not its question. The current opening sentence (\"Why do some states succeed at industrial policy while others fail under similar external conditions?\") is fine for a paper but slow for a card. Rewrite as: \"Upstream targeting (capital goods, heavy industry) generates stronger backward linkages and technological learning than downstream targeting, but requires higher state capacity to execute because these sectors face longer payoff horizons and greater rent-seeking pressures. The paper develops this argument with a comparative-historical account of five waves of industrial policy and contrasts Latin America's downstream trap under import substitution with East Asia's upstream-first heavy and chemical industrialization.\"")
P("The Subnational Industrial Policy Observatory does not belong on the Research page. It is a data platform, not an argument. Move it to Data, where it pairs naturally with the replication dataset and the Atlas. The Research page then carries three working papers plus the book, all argumentative.")
P("Add the year to the published article (\"Perspectives on Politics, 2025\" or whichever volume). A reader scanning the page should not have to click through to learn that.")
P("Taming Nature is the strongest single piece of intellectual real estate on the site. It deserves more than one paragraph. Consider a short standalone block above the working papers, with a one-line tagline and the four-country scope made visible as a small caption (Chile, Argentina, Brazil, United States; archival, comparative-historical, GIS).")

H("Policy versus Media: sharpen the boundary.", level=2)
P("\"Chile, Back to the Future\" (COHA) is filed under Policy; \"The Coming Copper Boom\" (Foreign Policy) is filed under Media. A reader cannot infer the rule from the examples. A defensible rule is: Policy = institutional outlet (think tank, lab, NGO, multilateral); Media = popular press (newspaper, magazine). Under that rule, COHA, NZIPL, Global Americans, and UNDP are Policy; Foreign Policy, The New Republic, La Jornada, and Phenomenal World are Media. Apply consistently or merge the two pages under \"Writing\" with two subheaders.")

H("Data page is too thin.", level=2)
P("Currently a single Dataverse link. Move the Observatory here, add the Atlas (even as a forward pointer), and consider listing the underlying Chile panel separately from the Perspectives on Politics replication package: the panel is reused across the Landholding paper and the book, and other researchers will want it independent of the published article.")

H("Atlas page: show, do not promise.", level=2)
P("The current page is a teaser card. Until the Atlas is online, the page should at minimum show one or two static maps from the panel (the 1862 department map and a fiscal-capacity choropleth would do), a list of the ten topics by name rather than as a count, and an indicative source list. The status badge is appropriate; the rest reads as marketing copy without an underlying product.")

H("Add a Talks or News section, even minimal.", level=2)
P("Three to five lines of recent activity (talks, workshops, fellowships, conferences) help readers understand whether the site is current. The CV already lists these; surfacing the last three or four on the homepage or as a small \"Recent\" block on Research costs little and signals that the site is maintained.")

# DESIGN
H("Design notes")

H("Typography.", level=2)
P("tokens.css declares Fraunces as --display but the Google Fonts request loads only Spectral and Manrope. The font-variation-settings on .serif and .overline (opsz, SOFT) are Fraunces axes and do nothing on Spectral. Either load Fraunces and use it for the display headings on page heroes (it would suit the editorial aesthetic well), or drop the unused tokens. The SVG favicon also uses Georgia rather than Spectral; a one-line fix to match the system.")

H("Link affordance.", level=2)
P("Default underlines use --rule-strong (#C9C0AC) which is barely visible against the warm paper. Hover darkens to the accent. This is on the edge of acceptable: a user scanning the page may miss in-line links, especially in the about-text block. Either darken the default underline to --ink-faint, or set text-decoration-thickness: 1.5px to compensate.")

H("Page-hero kicker duplicates the H1.", level=2)
P("Research / Research, Policy Analysis / Policy Analysis, Media / Media. Drop the kicker on pages where it repeats, or repurpose it for a one-line abstract (\"publications and working papers\" for Research, \"essays and op-eds\" for Media). The Atlas page does this well by using the kicker to repeat the page title and then carrying a description in the card; the others do not.")

H("The ornamental rule on the home page.", level=2)
P("The three dot-bullets in a hairline frame are pleasant but read as a Squarespace flourish in an otherwise restrained editorial system. A plain horizontal rule, or a single typographic mark (an asterism or a sectional pilcrow), would sit better with Spectral and the warm paper. Minor; subjective.")

H("The Cicarelli painting is cropped.", level=2)
P("painting-section uses object-fit: cover with object-position: center 40%, which cuts the foreground figures and the sky. The painting is a landscape with strong horizontal composition; a contain crop with a thin matted border would honor it. If the visual budget for height is tight, reduce max-height to 320px and let the full painting render.")

H("Photo size.", level=2)
P("The home photo is pulled at full resolution from the Squarespace CDN and rendered inside a 320px column. Serve a sized image (640px wide, 2x for retina) to cut homepage weight by an order of magnitude.")

# SEO / ACCESSIBILITY
H("SEO and accessibility")
P("Each page carries a meta description, which is good. None carry Open Graph or Twitter card tags, which means link previews on social media will show the favicon only. Adding og:title, og:description, og:image, and og:url to each page is a fifteen-minute task and substantially improves how the site looks when shared. A canonical link tag per page prevents duplicate-content concerns once the Spanish mirror is indexed.")
P("Two accessibility fixes worth making: add a \"Skip to main content\" link as the first focusable element on each page, and add aria-label=\"Primary\" to the injected nav element. Once nav is in static HTML these are one-line additions.")

# QUICK PRIORITY LIST
H("Priority order")
P("If you publish in the next week, do this in order: bake nav into HTML, self-host the two homepage images, delete font-compare.html, consolidate the CV files and drop the iframe, add the year to the Perspectives on Politics card, fix the Spanish title language mismatch, and rewrite the homepage opening. Everything else (link contrast, photo sizing, OG tags, Talks block, Atlas content, Observatory move) can wait until the next iteration.")

# IN THIS FOLDER
H("Proposed files in this folder")
P("Three revised pages are included as concrete examples of the recommendations above: index.html (new opening, local image references, static header), research.html (Landholding-style cards for the other three papers, Observatory removed, Perspectives year added), and cv.html (download-only, no iframe). The originals are not modified.")

out = "/Users/maximilianovejares/OneDrive - Johns Hopkins/Mis cosas/Projects/website/proposal/Vejares, Maximiliano. 2026. Website Critique.docx"
doc.save(out)
print(out)
